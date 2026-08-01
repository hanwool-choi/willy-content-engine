from datetime import date, timedelta
from pathlib import Path

import pytest
from docx import Document

from willy.models import DayWeather, Gender, LookAnalysis
from willy.publisher.folders import PUBLISH_STEM, REF_STEM, publish


def look(look_id: str, gender=Gender.MEN) -> LookAnalysis:
    return LookAnalysis(
        look_id=look_id, source="musinsa_snap", gender=gender, sleeve="short",
        outer=None, layers=1,
        fabric_weight="light", coverage="mid", temp_range=(24, 30), rain_ok=False,
        season="summer", style_tags=["미니멀"], palette=["ecru"],
    )


def week_of(days: int = 7) -> list[DayWeather]:
    out = []
    for i in range(days):
        d = date(2026, 8, 3) + timedelta(days=i)
        out.append(
            DayWeather(
                date=d, weekday_ko="월화수목금토일"[d.weekday()], temp_max=29,
                temp_min=24, precip_prob=10, sky="맑음", resolution="detailed",
            )
        )
    return out


@pytest.fixture
def setup(tmp_path: Path):
    """원본과 생성물 파일을 준비하고 (assignment, generated, week, outputs_root)를 돌려준다."""
    src_dir = tmp_path / "src"
    gen_dir = tmp_path / "gen"
    src_dir.mkdir()
    gen_dir.mkdir()

    week = week_of()
    assignment = {}
    generated = {}

    for day in week:
        for gender in (Gender.MEN, Gender.WOMEN):
            lid = f"{day.date.day}-{gender.value}"
            analysis = look(lid, gender)
            ref = src_dir / f"{lid}.jpg"
            ref.write_bytes(b"\xff\xd8ref")  # JPEG 매직 바이트
            analysis.image_path = ref

            gen = gen_dir / f"{lid}.png"
            # PNG 매직 바이트는 8바이트 전부 있어야 sniff()가 인식한다.
            gen.write_bytes(b"\x89PNG\r\n\x1a\n" + b"gen")

            assignment[(day.date, gender)] = analysis
            generated[(day.date, gender)] = gen

    return assignment, generated, week, tmp_path / "outputs"


def test_publish_creates_iso_week_folder(setup):
    assignment, generated, week, out = setup
    root = publish(assignment, week, generated, output_root=out)

    # 2026-08-03은 월요일, 그 주 목요일은 08-06 -> 8월 1주차
    assert root.name == "2026-08_W1"


def test_publish_creates_one_folder_per_day_with_weather_in_name(setup):
    assignment, generated, week, out = setup
    root = publish(assignment, week, generated, output_root=out)

    day_dirs = sorted(p.name for p in root.iterdir() if p.is_dir())
    assert len(day_dirs) == 7
    assert "08-03_월_맑음_29-24℃" in day_dirs


def test_publish_writes_ref_and_published_images(setup):
    assignment, generated, week, out = setup
    root = publish(assignment, week, generated, output_root=out)

    men_dir = root / "08-03_월_맑음_29-24℃" / "men"
    # ref는 JPEG 바이트로 준비했으니 .jpg, 생성물은 PNG 바이트니 .png로 남아야 한다.
    assert (men_dir / f"{REF_STEM}.jpg").exists()
    assert (men_dir / f"{PUBLISH_STEM}.png").exists()
    assert (men_dir / "analysis.json").exists()


def test_ref_filename_marks_do_not_publish(setup):
    """원본이 실수로 발행되지 않도록 파일명(스템)에 표식이 있어야 한다."""
    assert "발행금지" in REF_STEM


def test_publish_writes_item_doc_per_day(setup):
    assignment, generated, week, out = setup
    root = publish(assignment, week, generated, output_root=out)

    doc_path = root / "08-03_월_맑음_29-24℃" / "아이템정보.docx"
    assert doc_path.exists()

    doc = Document(str(doc_path))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "08-03" in text


def test_item_doc_created_even_without_metadata(setup):
    """메타데이터는 옵션이다. 없어도 문서는 만든다."""
    assignment, generated, week, out = setup
    root = publish(assignment, week, generated, output_root=out)

    doc = Document(str(root / "08-03_월_맑음_29-24℃" / "아이템정보.docx"))
    assert len(doc.tables) >= 1


def test_publish_writes_week_summary(setup):
    assignment, generated, week, out = setup
    root = publish(assignment, week, generated, output_root=out)

    assert (root / "_주간요약.docx").exists()


def test_publish_skips_empty_slot_without_crashing(setup):
    assignment, generated, week, out = setup
    empty_key = (week[0].date, Gender.MEN)
    assignment[empty_key] = None
    generated.pop(empty_key)

    root = publish(assignment, week, generated, output_root=out)

    assert not (root / "08-03_월_맑음_29-24℃" / "men").exists()
    assert (root / "08-03_월_맑음_29-24℃" / "women").exists()


def test_publish_preserves_generated_image_format(setup):
    """생성물이 PNG면 발행용도 .png다. 확장자와 내용이 어긋나면 안 된다."""
    assignment, generated, week, out = setup
    key = (week[0].date, Gender.MEN)
    generated[key].write_bytes(b"\x89PNG\r\n\x1a\n" + b"body")

    root = publish(assignment, week, generated, output_root=out)

    men_dir = root / week[0].folder_name / "men"
    assert (men_dir / f"{PUBLISH_STEM}.png").exists()
    assert not (men_dir / f"{PUBLISH_STEM}.jpg").exists()


def test_publish_preserves_generated_jpeg_format(setup):
    """생성물이 JPEG면 발행용은 .jpg여야 한다.

    위 PNG 테스트만으로는 확장자를 '.png'로 하드코딩해도 통과해버린다(픽스처
    기본값이 PNG이므로). 실제로 포맷을 판별하는지 검증하려면 PNG가 아닌
    포맷도 함께 확인해야 한다.
    """
    assignment, generated, week, out = setup
    key = (week[0].date, Gender.WOMEN)
    generated[key].write_bytes(b"\xff\xd8" + b"jpeg-body")

    root = publish(assignment, week, generated, output_root=out)

    women_dir = root / week[0].folder_name / "women"
    assert (women_dir / f"{PUBLISH_STEM}.jpg").exists()
    assert not (women_dir / f"{PUBLISH_STEM}.png").exists()


def test_publish_survives_corrupt_generated_image(setup):
    """최종 컨펌 이후 단계다. 이미지 한 장이 깨져도 주 전체를 날리지 않는다."""
    assignment, generated, week, out = setup
    generated[(week[0].date, Gender.MEN)].write_bytes(b"not an image at all")

    root = publish(assignment, week, generated, output_root=out)

    men_dir = root / week[0].folder_name / "men"
    assert not list(men_dir.glob(f"{PUBLISH_STEM}.*"))  # 발행용은 못 만든다
    assert (men_dir / "analysis.json").exists()          # 나머지는 정상
    assert (root / "_주간요약.docx").exists()
    # 다른 요일은 영향을 받지 않는다
    assert list((root / week[1].folder_name / "men").glob(f"{PUBLISH_STEM}.*"))


def test_publish_survives_corrupt_reference_image(setup):
    assignment, generated, week, out = setup
    assignment[(week[0].date, Gender.MEN)].image_path.write_bytes(b"")

    root = publish(assignment, week, generated, output_root=out)

    men_dir = root / week[0].folder_name / "men"
    assert not list(men_dir.glob(f"{REF_STEM}.*"))
    assert list(men_dir.glob(f"{PUBLISH_STEM}.*"))  # 발행용은 여전히 만들어진다
