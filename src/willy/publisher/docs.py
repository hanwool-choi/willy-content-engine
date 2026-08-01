"""워드 문서 생성. 메타데이터는 옵션이므로 없어도 문서는 만든다."""
from __future__ import annotations

from pathlib import Path

from docx import Document

from willy.models import Assignment, DayWeather, Gender

COLUMNS = ["제품명", "브랜드", "가격", "구매링크"]


def write_item_doc(path: Path, day: DayWeather, entries: dict[Gender, list[dict]]) -> None:
    """요일별 아이템 정보 문서.

    entries가 비어 있어도 표 골격은 남긴다. 사장님이 손으로 채울 수 있게.
    """
    doc = Document()
    doc.add_heading(
        f"{day.date.month:02d}-{day.date.day:02d} ({day.weekday_ko}) "
        f"{day.sky} {day.temp_max}/{day.temp_min}℃",
        level=1,
    )

    for gender in (Gender.MEN, Gender.WOMEN):
        doc.add_heading("남성" if gender is Gender.MEN else "여성", level=2)
        rows = entries.get(gender, [])

        table = doc.add_table(rows=1, cols=len(COLUMNS))
        table.style = "Table Grid"
        for i, name in enumerate(COLUMNS):
            table.rows[0].cells[i].text = name

        if not rows:
            blank = table.add_row().cells
            blank[0].text = "(수집된 아이템 정보 없음)"
        for row in rows:
            cells = table.add_row().cells
            cells[0].text = str(row.get("name", ""))
            cells[1].text = str(row.get("brand", ""))
            cells[2].text = str(row.get("price", ""))
            cells[3].text = str(row.get("url", ""))

    doc.save(str(path))


def write_week_summary(path: Path, week: list[DayWeather], assignment: Assignment) -> None:
    doc = Document()
    doc.add_heading("이번주 [내일 뭐입지?] 요약", level=1)

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    for i, name in enumerate(["날짜", "요일", "날씨", "기온", "배정 상태"]):
        table.rows[0].cells[i].text = name

    for day in week:
        men = assignment.get((day.date, Gender.MEN))
        women = assignment.get((day.date, Gender.WOMEN))
        filled = sum(1 for x in (men, women) if x is not None)

        cells = table.add_row().cells
        cells[0].text = day.date.isoformat()
        cells[1].text = day.weekday_ko
        cells[2].text = day.sky
        cells[3].text = f"{day.temp_max}/{day.temp_min}℃"
        cells[4].text = f"{filled}/2"

    doc.save(str(path))
